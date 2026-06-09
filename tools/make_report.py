from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "report" / "ElectraLens_AI_결과보고서.docx"
SCREENSHOTS = ROOT / "screenshots"
WEB_EXAMPLES = ROOT / "web_examples"


BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(18, 36, 46)
MUTED = RGBColor(83, 100, 108)
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"


def set_font(run, size=None, bold=None, color=None, name="Malgun Gothic"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, fill=None) -> None:
    if fill:
        shade_cell(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, size=10.5, bold=bold, color=DARK)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.8)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Title", 24, DARK),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK),
    ]:
        style = doc.styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    run = p.add_run("ElectraLens AI 결과보고서")
    set_font(run, size=24, bold=True, color=DARK)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    run = sub.add_run("사진 한 장으로 전기전자 상황을 읽는 오프라인 데스크톱 AI 조교")
    set_font(run, size=12.5, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.35)
    table.columns[1].width = Inches(5.25)
    rows = [
        ("과제 주제", "AI 코딩 활용을 통한 전기전자 응용 프로그램 개발"),
        ("개발 형태", "Windows 단일 실행파일 기반 데스크톱 프로그램"),
        ("핵심 가치", "텍스트 입력 없이 이미지 열기, 클립보드, 샘플 버튼으로 전기전자 상황 분석"),
        ("제출 구성", "실행파일, 소스 코드, 원본 DOCX 보고서"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_text(row.cells[0], label, bold=True, fill=PALE_BLUE)
        set_cell_text(row.cells[1], value)

    doc.add_paragraph()


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10.8, color=DARK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, size=10.5, color=DARK)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, size=10.5, color=DARK)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=9, color=MUTED)


def add_picture(doc: Document, filename: str, caption: str, width=6.3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(SCREENSHOTS / filename), width=Inches(width))
    add_caption(doc, caption)


def add_abs_picture(doc: Document, path: Path, caption: str, width=6.3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_feature_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [1.4, 2.4, 2.7]
    headers = ["기능", "일반 사용자 가치", "공대생 활용성"]
    for i, text in enumerate(headers):
        table.columns[i].width = Inches(widths[i])
        set_cell_text(table.rows[0].cells[i], text, bold=True, fill=LIGHT)
    rows = [
        ("멀티탭 안전 분석", "과부하와 발열 위험을 직관적으로 확인", "P=VI, I²R, 정격전류 개념 연결"),
        ("LED 회로 해석", "회로 사진을 보고 위험/정상 여부 이해", "옴의 법칙, 전력손실, 부품 정격 학습"),
        ("스코프 파형 분석", "파형 사진에서 주기와 진폭 의미 파악", "T, f, Vpp, Vrms 계측 개념 학습"),
        ("어댑터 라벨 판독", "충전기 호환성 판단 보조", "AC/DC, 출력 정격, 극성 개념 확인"),
        ("스마트 프로브", "사진 위 특정 구역을 드래그해 빠른 설명 확인", "부품, 배선, 파형선, 라벨 영역을 국소적으로 관찰"),
    ]
    for row_data in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_data):
            set_cell_text(cell, text)


def add_file_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.3)
    set_cell_text(table.rows[0].cells[0], "파일", bold=True, fill=LIGHT)
    set_cell_text(table.rows[0].cells[1], "역할", bold=True, fill=LIGHT)
    rows = [
        ("dist/ElectraLensAI.exe", "제출/실행용 단일 Windows 실행파일"),
        ("electralens/analyzer.py", "모드 분류 + 모드별 전기전자 해석 엔진(UI 비의존)"),
        ("electralens/features.py, waveform.py", "이미지 특징 추출, 파형 추적 + 격자 검출"),
        ("electralens/app.py, cli.py", "Tkinter GUI, 명령행/헤드리스 진입점"),
        ("build_exe.ps1", "PyInstaller 기반 exe 빌드 스크립트"),
        ("README.md, docs/DESIGN.md", "사용/빌드 설명, 설계·변경 내역"),
        ("tests/", "pytest 단위 테스트"),
    ]
    for row_data in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_data):
            set_cell_text(cell, text)


def add_web_source_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    widths = [1.25, 2.05, 1.45, 1.75]
    headers = ["분석 모드", "이미지 출처", "라이선스", "파일명"]
    for i, text in enumerate(headers):
        table.columns[i].width = Inches(widths[i])
        set_cell_text(table.rows[0].cells[i], text, bold=True, fill=LIGHT)
    rows = [
        ("안전", "Wikimedia Commons, NEMA-AC-Power-Plugs.jpg, Evan-Amos", "Public domain", "web_safety_NEMA_AC_Power_Plugs.jpg"),
        ("회로", "Wikimedia Commons, Solderless Breadboard with LEDs.jpg, Ilikefood", "Copyrighted free use", "web_circuit_Solderless_Breadboard_with_LEDs.jpg"),
        ("파형", "NIST Image Gallery, Oscilloscope traces, Hamilton/NIST", "Public domain / U.S. gov work", "web_scope_NIST_oscilloscope_traces.jpg"),
    ]
    for row_data in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_data):
            set_cell_text(cell, text)


def build() -> None:
    doc = Document()
    style_doc(doc)
    add_title_block(doc)

    doc.add_heading("1. 프로그램 소개", level=1)
    add_para(
        doc,
        "ElectraLens AI는 전기전자 상황을 사진 기반으로 해석하는 데스크톱 프로그램이다. "
        "사용자는 제품명, 전압, 전류를 일일이 입력하지 않아도 되며, 이미지 파일을 열거나 클립보드 이미지를 가져오고 분석 모드를 마우스로 선택하면 된다. "
        "이미지 열기, 클립보드 가져오기, 샘플 버튼을 중심으로 사용할 수 있고, 이미지 파일을 실행파일 아이콘 위에 올려 바로 열 수도 있다. 사진 위의 궁금한 구역을 드래그하면 스마트 프로브가 해당 주변 영역만 따로 해석한다. "
        "프로그램은 멀티탭 안전, 회로 해석, 오실로스코프 파형, 어댑터 라벨을 전기전자 공식과 연결해 설명한다.",
    )
    add_para(
        doc,
        "기존의 전공 공식 계산기는 사용자가 숫자를 알고 직접 입력해야 한다는 한계가 있다. 이 프로그램은 현실 세계의 사진을 먼저 받아들이고, 그 사진에서 가능한 해석 방향을 제시한다는 점에서 일반인과 공대생 모두에게 더 자연스럽다.",
    )
    add_picture(doc, "electralens_main.png", "그림 1. ElectraLens AI 실행 화면")

    doc.add_heading("2. 개발 목표와 차별점", level=1)
    add_bullets(
        doc,
        [
            "웹사이트가 아니라 실행파일 하나를 더블클릭하면 켜지는 프로그램으로 제작한다.",
            "사용자가 텍스트를 거의 입력하지 않아도 마우스와 이미지 중심으로 사용할 수 있게 한다.",
            "사진 전체 분석뿐 아니라 사용자가 드래그로 지정한 구역만 빠르게 읽는 스마트 프로브를 제공한다.",
            "인터넷이나 API 키가 없어도 오류 없이 기본 분석이 가능하도록 오프라인 엔진을 포함한다.",
            "일반인에게는 안전과 호환성 판단을, 공대생에게는 전공 공식과 계측 개념을 함께 제공한다.",
            "보고서 제출을 위해 실행 화면, 개발 과정, 소스 구조, 느낀 점이 모두 드러나도록 정리한다.",
        ],
    )

    doc.add_heading("3. 핵심 기능", level=1)
    add_feature_table(doc)
    doc.add_paragraph()
    add_picture(doc, "electralens_scope.png", "그림 2. 오실로스코프 파형 분석 예시")
    add_picture(doc, "electralens_label.png", "그림 3. 어댑터 라벨 분석 예시")

    doc.add_heading("4. 실제 웹 이미지 분석 결과", level=1)
    add_para(
        doc,
        "기본 샘플만 사용하면 프로그램이 미리 정해진 그림에만 동작하는 것처럼 보일 수 있다. 그래서 Wikimedia Commons와 NIST에서 공개된 실제 이미지를 내려받아 ElectraLens AI에 입력하고, 같은 분석 엔진으로 결과 화면을 생성했다.",
    )
    add_web_source_table(doc)
    doc.add_paragraph()
    add_abs_picture(doc, WEB_EXAMPLES / "results" / "real_safety_result.png", "그림 4. 실제 웹 이미지: 전원 플러그 사진 안전 분석 결과", width=6.3)
    add_abs_picture(doc, WEB_EXAMPLES / "results" / "real_circuit_result.png", "그림 5. 실제 웹 이미지: 브레드보드 LED 회로 분석 결과", width=6.3)
    add_abs_picture(doc, WEB_EXAMPLES / "results" / "real_scope_result.png", "그림 6. 실제 웹 이미지: NIST 오실로스코프 파형 분석 결과", width=6.3)
    add_para(
        doc,
        "이 실제 이미지들은 별도 압축파일로 함께 제출할 수 있도록 submission/ElectraLensAI_Web_Images.zip에 묶었다. 따라서 보고서에 들어간 실행 결과가 어떤 원본 이미지에서 나왔는지 다시 확인할 수 있다.",
    )

    doc.add_heading("5. AI 활용 과정", level=1)
    add_para(
        doc,
        "처음에는 전력 소비나 멀티탭 안전을 직접 입력해 계산하는 앱을 구상했지만, 사용자가 귀찮음을 크게 느낀다는 조건을 추가하면서 방향을 바꾸었다. "
        "AI와 대화하며 '교수급 지능을 가진 사용자가 텍스트 입력을 싫어한다면 어떤 앱이 실용적인가'를 기준으로 다시 설계했고, 결과적으로 이미지 중심 앱이라는 방향이 나왔다.",
    )
    add_numbered(
        doc,
        [
            "아이디어 발산: 단순 공식 계산기, 전기요금 계산기, 회로 도우미 등을 비교했다.",
            "요구사항 정제: 웹사이트 제외, 단일 exe, 마우스 중심, 오프라인 안정성을 핵심 조건으로 확정했다.",
            "구현 보조: AI 코딩 도구를 활용해 Tkinter GUI, 샘플 이미지 자동 생성, 이미지 특징 추출, 스마트 프로브, 분석 문장 생성 로직을 빠르게 작성했다.",
            "검증과 수정: 샌드박스 저장 권한 문제, exe 빌드 도구 부재, 실행파일 검증 문제를 발견하고 예외 처리와 빌드 과정을 수정했다.",
            "보고서 작성: 개발 과정과 실행 화면을 원본 DOCX 형식으로 정리했다.",
        ],
    )
    add_para(
        doc,
        "앱 내부의 AI는 인터넷 연결이 필요한 대형 모델에 전적으로 의존하지 않는다. 이미지의 색상, 에지 밀도, 파형 후보점, 파일/샘플 맥락을 바탕으로 분석 모드를 분류하고, 전기전자 지식 베이스를 이용해 설명과 행동 권고를 생성한다. API 기반 비전 모델을 붙일 수도 있지만, 제출용 실행파일은 다른 사람 PC에서도 안정적으로 켜져야 하므로 기본 기능은 오프라인으로 설계했다.",
    )

    doc.add_heading("6. 구현 구조", level=1)
    add_file_table(doc)
    doc.add_paragraph()
    add_para(
        doc,
        "프로그램은 Python의 Tkinter로 GUI를 만들고, Pillow와 NumPy로 이미지를 처리한다. PyInstaller를 사용해 Python 런타임과 필요한 라이브러리를 하나의 실행파일 안에 포함했다. 샘플 이미지는 별도 파일을 배포하지 않아도 되도록 프로그램 실행 중 자동 생성된다.",
    )

    doc.add_heading("7. 테스트 및 배포 안정성", level=1)
    add_bullets(
        doc,
        [
            "샘플 4종 안전/회로/파형/라벨 분석이 정상적으로 분류되는지 smoke test를 수행했다.",
            "PyInstaller one-file, windowed 옵션으로 dist/ElectraLensAI.exe를 생성했다.",
            "exe 자체에서 스크린샷 생성 명령을 실행해 패키징 후에도 이미지 생성과 분석 로직이 동작함을 확인했다.",
            "이미지 파일을 exe로 전달하는 실행 인자를 처리해, 이미지 파일을 실행파일에 드래그해도 바로 열리도록 했다.",
            "GUI를 실제로 실행해 3초 이상 정상 유지되는지 확인했다.",
            "저장 권한이 없는 환경에서도 AppData, 임시 폴더, 현재 폴더 순서로 자동 우회하도록 설계했다.",
        ],
    )

    doc.add_heading("8. 한계와 개선 방향", level=1)
    add_para(
        doc,
        "현재 버전은 완전한 OCR이나 실제 부품 인식 모델을 내장한 것은 아니다. 대신 오프라인 실행 안정성을 우선하여 보수적인 분석 엔진과 전기전자 지식 기반 설명을 제공한다. 향후 개선한다면 OpenAI Vision API나 로컬 OCR 모델을 선택 기능으로 연결해 실제 라벨 숫자와 회로 부품을 더 정확히 읽을 수 있다.",
    )
    add_para(
        doc,
        "또한 카메라 실시간 촬영, 부품 영역 자동 검출, 회로도 자동 Netlist 변환 기능을 추가하면 공대생 실험 수업에서 더욱 유용한 도구가 될 수 있다.",
    )

    doc.add_heading("9. 느낀 점", level=1)
    add_para(
        doc,
        "이번 과제를 하면서 AI 코딩은 단순히 코드를 대신 쓰게 하는 도구가 아니라, 요구사항을 계속 다듬고 구현 가능한 형태로 바꾸는 협업 도구라는 점을 느꼈다. 처음 생각한 프로그램은 입력 부담이 컸지만, AI와 대화하며 사용자의 귀찮음까지 설계 조건으로 넣자 더 좋은 방향이 나왔다.",
    )
    add_para(
        doc,
        "전기전자 지식도 공식 자체를 외우는 것보다, 현실의 사진과 문제 상황에 연결할 때 훨씬 실용적으로 느껴졌다. 앞으로 공학 문제를 해결할 때도 AI를 단순 검색 도구가 아니라 아이디어 설계, 구현, 검증을 함께하는 도구로 활용할 수 있겠다고 생각했다.",
    )

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "ElectraLens AI 결과보고서"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header.runs:
        set_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "AI 코딩 활용 전기전자 응용 프로그램 개발"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer.runs:
        set_font(footer.runs[0], size=9, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
